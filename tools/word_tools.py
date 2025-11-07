"""
Word Tools - DOCX parsing, extraction, and analysis tools.
Handles Word document processing for batch records, procedures, and reports.
"""

import logging
import re
from pathlib import Path
from typing import Dict, Any, List, Optional

logger = logging.getLogger(__name__)

from docx import Document


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text content from Word document.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    logger.info(f"Extracting text from DOCX: {docx_path}")
    
    try:
        path = Path(docx_path)
        if not path.exists():
            return f"Error: File not found: {docx_path}"
        
        doc = Document(docx_path)
        text_content = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        return '\n'.join(text_content)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return f"Error: {str(e)}"


def extract_tables_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """
    Extract tables from Word document.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        List of extracted tables with structured data
    """
    logger.info(f"Extracting tables from DOCX: {docx_path}")
    
    try:
        path = Path(docx_path)
        if not path.exists():
            return [{"error": f"File not found: {docx_path}"}]
        
        doc = Document(docx_path)
        all_tables = []
        
        for table_idx, table in enumerate(doc.tables, 1):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            all_tables.append({
                "table_id": f"table_{table_idx}",
                "data": table_data,
                "rows": len(table_data),
                "columns": len(table_data[0]) if table_data else 0
            })
        
        return all_tables
    except Exception as e:
        logger.error(f"Error extracting tables from DOCX: {e}")
        return [{"error": str(e)}]


def parse_bmr_docx(docx_path: str) -> Dict[str, Any]:
    """
    Parse Batch Manufacturing Record (BMR) from Word document.
    
    Args:
        docx_path: Path to BMR DOCX file
        
    Returns:
        Structured BMR data with batch info, manufacturing steps
    """
    logger.info(f"Parsing BMR from DOCX: {docx_path}")
    
    try:
        text = extract_text_from_docx(docx_path)
        tables = extract_tables_from_docx(docx_path)
        
        bmr_data = {
            "document_type": "Batch Manufacturing Record (BMR)",
            "source": docx_path,
            "raw_text": text,
            "tables": tables,
            "manufacturing_steps": [],
            "metadata": extract_metadata_from_docx(docx_path)
        }
        
        # Extract batch number
        batch_patterns = [
            r"Batch\s*(?:Number|No\.?|#)?\s*:?\s*([A-Z0-9\-]+)",
            r"BMR\s*(?:Number|No\.?|#)?\s*:?\s*([A-Z0-9\-]+)"
        ]
        for pattern in batch_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                bmr_data["batch_number"] = match.group(1)
                break
        
        # Extract product name
        product_patterns = [
            r"Product\s*(?:Name)?\s*:?\s*([A-Za-z0-9\s\-]+)",
            r"Material\s*:?\s*([A-Za-z0-9\s\-]+)"
        ]
        for pattern in product_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                bmr_data["product_name"] = match.group(1).strip()
                break
        
        # Parse manufacturing steps from tables
        for table_info in tables:
            table_data = table_info.get("data", [])
            if table_data and len(table_data) > 1:
                headers = table_data[0]
                for row in table_data[1:]:
                    if len(row) == len(headers):
                        step_entry = dict(zip(headers, row))
                        bmr_data["manufacturing_steps"].append(step_entry)
        
        return bmr_data
        
    except Exception as e:
        logger.error(f"Error parsing BMR DOCX: {e}")
        return {
            "document_type": "BMR",
            "error": str(e),
            "source": docx_path
        }


def parse_sop_docx(docx_path: str) -> Dict[str, Any]:
    """
    Parse Standard Operating Procedure (SOP) from Word document.
    
    Args:
        docx_path: Path to SOP DOCX file
        
    Returns:
        Structured SOP data with SOP number, title, procedures
    """
    logger.info(f"Parsing SOP from DOCX: {docx_path}")
    
    try:
        text = extract_text_from_docx(docx_path)
        tables = extract_tables_from_docx(docx_path)
        
        sop_data = {
            "document_type": "Standard Operating Procedure (SOP)",
            "source": docx_path,
            "raw_text": text,
            "tables": tables,
            "procedures": [],
            "metadata": extract_metadata_from_docx(docx_path)
        }
        
        # Extract SOP number
        sop_patterns = [
            r"SOP\s*(?:Number|No\.?|#)?\s*:?\s*([A-Z0-9\-]+)",
            r"Document\s*(?:Number|No\.?)?\s*:?\s*([A-Z0-9\-]+)"
        ]
        for pattern in sop_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                sop_data["sop_number"] = match.group(1)
                break
        
        # Extract title
        title_patterns = [
            r"Title\s*:?\s*([^\n]+)",
            r"Procedure\s*(?:Name|Title)?\s*:?\s*([^\n]+)"
        ]
        for pattern in title_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                sop_data["title"] = match.group(1).strip()
                break
        
        # Extract effective date
        date_patterns = [
            r"Effective\s*Date\s*:?\s*([0-9\/\-]+)",
            r"Revision\s*Date\s*:?\s*([0-9\/\-]+)"
        ]
        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                sop_data["effective_date"] = match.group(1)
                break
        
        # Parse procedure steps from numbered lists
        procedure_pattern = r"(\d+\.?\d*)\s+([^\n]+)"
        procedure_matches = re.findall(procedure_pattern, text)
        for step_num, step_text in procedure_matches:
            if len(step_text) > 10:  # Filter out short matches
                sop_data["procedures"].append({
                    "step": step_num,
                    "description": step_text.strip()
                })
        
        return sop_data
        
    except Exception as e:
        logger.error(f"Error parsing SOP DOCX: {e}")
        return {
            "document_type": "SOP",
            "error": str(e),
            "source": docx_path
        }


def extract_metadata_from_docx(docx_path: str) -> Dict[str, Any]:
    """
    Extract metadata from Word document.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Dictionary with DOCX metadata
    """
    logger.info(f"Extracting metadata from DOCX: {docx_path}")
    
    try:
        path = Path(docx_path)
        if path.exists():
            stat = path.stat()
            metadata = {
                "filename": path.name,
                "size_bytes": stat.st_size,
                "size_kb": round(stat.st_size / 1024, 2),
                "path": str(path),
                "exists": True
            }
            
            # Extract DOCX-specific metadata
            try:
                doc = Document(docx_path)
                core_props = doc.core_properties
                metadata["docx_metadata"] = {
                    "author": core_props.author,
                    "created": str(core_props.created) if core_props.created else None,
                    "modified": str(core_props.modified) if core_props.modified else None,
                    "title": core_props.title,
                    "subject": core_props.subject
                }
                metadata["num_paragraphs"] = len(doc.paragraphs)
                metadata["num_tables"] = len(doc.tables)
            except Exception as e:
                logger.error(f"Error reading DOCX metadata: {e}")
            
            return metadata
        return {"filename": path.name, "exists": False}
    except Exception as e:
        logger.error(f"Error extracting metadata from DOCX: {e}")
        return {"error": str(e)}

